def seal_estimation(road_length,ground_type,culvert_number,culvert_diameter,culvert_price,citizen,citizen_request,land_number,basin_name,basin_number,area):
    if ground_type == "طبقتين":
        subgrade_width=6
        subgrade_price=1
        topping_width=5.50
        topping_price=2.25
        basecourse_width=4.5
        basecourse_price=2.25
        mc_price=0.5
        rc_price=1.00
        line_layers="طبقة "
    elif ground_type=="طبقة" :
        subgrade_width=6
        subgrade_price=1.25
        topping_width=0
        topping_price=0
        basecourse_width=4.5
        basecourse_price=2.50
        mc_price=0.5
        rc_price=1.00
        line_layers="طبقة فرشيات"

    #----------------------------------------------------------------------------------------------------
    culvert_number=int(culvert_number)
    culvert_price=int(culvert_price)
    #----------------------------------------------------------------------------------------------------

    detail={'width':[subgrade_width,topping_width,basecourse_width,basecourse_width,basecourse_width,basecourse_width],
            'price':[subgrade_price,topping_price,basecourse_price,mc_price,rc_price,rc_price]}
    #----------------------------------------------------------------------------------------------------
    import pandas as pd
    detail=pd.DataFrame(detail,index=['subgrade','Topping','Base course','MC','Rc1','Rc2'])
    detail['Road Length']=road_length
    #----------------------------------------------------------------------------------------------------
    detail['pavement price']=road_length*detail['price']*detail['width']
    #----------------------------------------------------------------------------------------------------
    detail = detail.rename(columns={'width': 'عرض الطبقة'})
    detail = detail.rename(columns={'price': 'السعر'})
    detail = detail.rename(columns={'Road Length': 'طول الطريق'})
    detail = detail.rename(columns={'pavement price': 'سعر الطبقة'})
    #----------------------------------------------------------------------------------------------------
    culvert_numberd='عدد العبارات'+ '(' +str(int(culvert_number))+')'
    culvert_diameterd='قطر العبارة'+ '(' +str(int(culvert_diameter))+')'
    culvert_priced='سعر العبارة'+ '(' +str(int(culvert_price))+')' + ' ' + 'دينار'
    detail.loc['culvert']=[culvert_diameterd,culvert_numberd, culvert_priced,culvert_number*culvert_price]
    #----------------------------------------------------------------------------------------------------

    cost_total=detail['سعر الطبقة'].sum()
    detail.loc['الكلفة الكلية للطريق ']=['--','--','--',cost_total]

    #----------------------------------------------------------------------------------------------------
    import num2words
    from num2words import num2words

    import math
    n=cost_total
    count=0
    while(n>0):
        count=count+1
        n=n//10
    x= math.floor(cost_total/(10**(count-2)))
    cost_total=(x+1)*10**(count-2)
    print ('الكلفة بعد التقريب ' + '(' + str(cost_total)+ ')' )
    detail.loc['الكلفة الكلية للطريق بعد التقريب ']=['--','--','--',cost_total]

    #----------------------------------------------------------------------------------------------------
    import numpy as np

    import pandas as pd
    import numpy as np
    from openpyxl import load_workbook
    citizen_name='طريق'+' ' +citizen

    path = "projectOne//static//Seal_Coat_Lists.xlsx"
    inf1=land_number
    inf2=basin_name
    inf3=basin_number
    inf4=area
    detail2=pd.DataFrame({'اسم المواطن':citizen_name,"اسم المنطقة":inf4,'اسم الحوض':inf2,'رقم الحوض':inf3,
                         'رقم القطعة':inf1},index=["-"])
    #----------------------------------------------------------------------------------------

    book = load_workbook(path)

    writer = pd.ExcelWriter(path , engine = 'openpyxl',kind='a')
    writer.book=book
    detail.to_excel(writer, sheet_name =citizen_name)
    detail2.to_excel(writer, sheet_name = citizen_name,startrow=14,index=False)


    writer.save()

    #----------------------------------------------------------------------------------------

    sheet= book.get_sheet_by_name("1")
    road_list= pd.read_excel("projectOne//static//Seal_Coat_Lists.xlsx", sheet_name="1")
    xx=road_list.last_valid_index()
    xx=xx+3
    value1='A'+str(xx)
    value2='B'+str(xx)
    value3='C'+str(xx)
    value4='D'+str(xx)
    value5='E'+str(xx)
    value6='F'+str(xx)
    sheet[value1] =citizen
    sheet[value2] =area
    sheet[value3] =inf2
    sheet[value4] =inf3
    sheet[value5] =inf1
    sheet[value6] =cost_total
    writer.save()

    #-----------------------------------------------------------------------------------------


    line1='   اشارة الى الطلب  المقدم من المواطن {one} '.format(one=citizen)
    line2= ' رقم ({two}) '.format(two=citizen_request)
    line3='و المتضمن فتح الطريق المرسم على مخططات الاراضي والمساحة بطول ({three}) '.format(three=road_length)
    line4= ' والواصل لقطعة الارض رقم ({four}) ' .format(four=land_number)
    line5= ' من حوض {five} '.format(five=basin_name)
    line6= ' رقم ({six}) '.format(six=basin_number)
    line7= ' في منطقة {seven} '.format(seven=area)

    line8=  '   نرجو عطوفتكم التكرم بمخاطبة رئيس واعضاء مجلس محافظة مادبا  بالموافقة على '
    line9= ' تخصيص مبلغ وقدره'
    line9l='('+str(cost_total)+')'
    line9t=  num2words( cost_total, lang='ar')
    line9e=' دينار'
    line9d="للعمل على تنفيذ {ten}".format(ten=line_layers)
    line10=' ووجهين {ten} '.format(ten='Seal Coat')
    line10e='للطريق اعلاه'
    line11=' ووضعه على موازنة السنوات القادمة وذلك لاجراء اللازم لخدمة مستخدمي الطريق وحسب الاصول مراعيين  كافة تعاميم الوزارة المتعلقة  بهذا الخصوص  '
    import docx
    from docx import Document
    book_to= '\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'عطوفة محافظ مادبا'
    book_subject='الموضوع:تخصيص مبلغ '
    book_paragraph1=line1+line2+line3+line4+line5+line6+line7+'.'
    book_paragraph2=line8+line9+line9l+line9t+line9e+line9d+line10+line10e+line11+'.'
    book_s1='وتفضلوا عطوفتكم قبول فائق الاحترام'
    book_s2= 'مدير اشغال محافظة مادبا ' + '\n' + ' المهندسة عروبه عويس  '
    book_s3=('نسخة / مدير ادارة شؤون المحافظــات'+'\n' + 'نسخه / مساعد المدير لشؤون الطـرق'
    +'\n'+'نسخه / ر.ق  قسم الطرق الزراعيــــة'+'\n'+'نسخه / للملــــــــــــــــــــــــــــــــــــف'
    +'\n'+': المرفقات  '+'\n'+ 'صوره عن الاستدعــــــــــاء ' + '\n' +'صورة عن تقرير المســـــاح '+'\n'+
    'صورة عن محطط الاراضي ')
    citizen_name='طريق'+' ' +citizen+'.docx'


    import docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

#--------------------------------------------------------------------------------------
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    sealcoat_text=Document()

#

#-----------------------------------------------------------------------
    # sealcoat_text.add_paragraph(book_to).alignment = WD_ALIGN_PARAGRAPH.CENTER
    #columns_removal.add_paragraph(book_to).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L1=sealcoat_text.add_paragraph()
    L1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner=L1.add_run(book_to)
    runner.bold=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True
#-----------------------------------------------------------------------------


    #columns_removal.add_paragraph(book_subject).alignment=WD_ALIGN_PARAGRAPH.RIGHT
#     sealcoat_text.add_paragraph(book_subject).alignment=WD_ALIGN_PARAGRAPH.RIGHT

    L2=sealcoat_text.add_paragraph()
    L2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner=L2.add_run(book_subject)
    # runner.font.BoldBi=True
    runner.bold=True
    runner.underline=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True



#-----------------------------------------------------------------------------------


#sealcoat_text.add_paragraph(book_paragraph1).alignment = WD_ALIGN_PARAGRAPH.LEFT

 #columns_removal.add_paragraph(book_paragraph1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    L3=sealcoat_text.add_paragraph()
    L3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L3.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    L3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    runner=L3.add_run(book_paragraph1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    runner.font.rtl=True

#------------------------------------------------------------------------------------

# sealcoat_text.add_paragraph(book_paragraph2).alignment = WD_ALIGN_PARAGRAPH.LEFT

    L4=sealcoat_text.add_paragraph()
    L4.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    L4.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    runner=L4.add_run(book_paragraph2)
    runner2=L4.add_run()
    runner.font.rtl=True
    runner.font.name="Arial"
    runner.font.size=Pt(15)


#------------------------------------------------------------------------------------
 #columns_removal.add_paragraph(book_s1).alignment = WD_ALIGN_PARAGRAPH.CENTER
 #sealcoat_text.add_paragraph(book_s1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L5=sealcoat_text.add_paragraph()
    L5.alignment=WD_ALIGN_PARAGRAPH.CENTER
    runner=L5.add_run(book_s1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    # runner.font.rtl=True

#------------------------------------------------------------------------------------
#sealcoat_text.add_paragraph(book_s2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    L6=sealcoat_text.add_paragraph()
    L6.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    runner=L6.add_run(book_s2)
    # runner.font.rtl=True
    runner.bold=True
    runner.font.name='Arial'
    runner.font.size=Pt(15)


#------------------------------------------------------------------------------------


# sealcoat_text.add_paragraph(book_s3).alignment = WD_ALIGN_PARAGRAPH.LEFT


    L7=sealcoat_text.add_paragraph()
    L7.alignment=WD_ALIGN_PARAGRAPH.LEFT
    runner=L7.add_run(book_s3)
    # runner.font.rtl=True
    runner.font.name='Arial'
    runner.font.size=Pt(14)

#------------------------------------------------------------------------------------

    docName="projectOne//static//"+ citizen+".docx"

    sealcoat_text.save(docName)
    return citizen
