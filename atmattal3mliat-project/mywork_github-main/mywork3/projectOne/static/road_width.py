def widthRequest_function(citizen,citizen_request,request_date,land_number,basin_name,basin_number
,area,road_length,road_width):

    #-------------------------------------------------------------------------------
    #الادخالات
    # citizen=input('ادخل اسم المواطن'+'\n') ok
    # citizen_name='طريق'+' '+citizen +'.' +'docx'
    # column_number=input(' ادخل عدد الاعمدة '+'\n') ok
    # citizen_request=input("ادخل رقم الطلب"+'\n') ok
    # land_number=input('ادخل رقم الارض'+'\n') ok
    # basin_name=input('ادخل اسم الحوض'+'\n') ok
    # basin_number=input('ادخل رقم الحوض'+'\n')  ok
    # area=input('ادخل اسم القرية او المنطقة'+'\n') ok
    # road_length=input('ادخل طول الطريق'+'\n')  ok
    #------------------------------------------------------------------------------------------------
    book_to='\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+ "عطوفة محافظ مأدبـــــــــا"
    book_subject= "الموضوع:- رفع سعه طريق"
    #------------------------------------------------------------------------------------------------
    #الفقرة الاولى
    line1='اشارة الى الطلب  المقدم من السيد {one} '.format(one=citizen)
    line2= ' رقم ({two}) '.format(two=citizen_request)
    line2t=' تاريخ {date} '.format(date=request_date)
    line3='بخصوص فتح الطريق المرسم على مخططات الاراضي والمساحة  '
    line4= ' والواصل لقطعة الارض رقم({four}) ' .format(four=land_number) +' '
    line5= 'حوض {five} '.format(five=basin_name)
    line6= ' رقم ({six}) '.format(six=basin_number)
    line7= ' من اراضي {seven} '.format(seven=area)
    line8= "والمرسم على مخططات دائرة الأراضي والمساحة بسعة ({eight})".format(eight=road_width)
    line9= " وبطول ({nine})".format(nine=road_length)
    line10='.'
    book_paragraph1="   "+line1+line2+line2t+line3+line4+line5+line6+line7+line8+line9+line10
    #----------------------------------------------------------------------------------------------------
    #الفقرة الثانية
    line11="   راجياً من عطوفتكم التكرم بالإيعاز لمديرية الاراضي والمساحة / مديرية هندسة البلديات "
    line13= "برفع سعة الطريق من ({width}) ".format(width=road_width)
    line14='الى السعة المطلوبة حسب النظام وذلك  لنتمكن من فتح وتعبيد الطريق حسب الأصول ,'
    line15=" ويكون التنسيق مع مندوبنا لهـــذه الغايــــة   المراقب / عبد الرحمن  0772641295" +'.'
    book_paragraph2="   "+line11+" "+ line13+line14+line15+'\n'
    #--------------------------------------------؛------------------------------------------------------------------
    book_s1='وتفضلوا عطوفتكم قبول فائق الاحترام'
    book_s2= 'مدير اشغال محافظة مادبا ' + '\n' + ' المهندسة عروبه عويس  ' +'\n'+'\n'
    book_s3= 'نسخه/مديـــــر دائرة الاراضي والمساحة'+'\n'+'نسخه/مديـــــر هندســـة البلـــديــــــــــات'+'\n'+'نسخه / ر.ق الطرق الزراعيــــــــــــــــة'+'\n'+'نسخه/للملـــــــــــــــــــــــــــــــــــــــــــف'
    book_s4="المرفقات:"+'\n'+'صوره عن تقرير  المساح المرخص'+'\n'+'صوره عن الطلب اعــــــــــــــــلاه'+'\n'+"صورة عن مخطط الاراضـــــــــي"
    #--------------------------------------------------------------------------------------------------------------
    book_note= "\n" + "برنامج اتمتة الخطابات الرسمية" +'\n'+'منشا البرنامج:م.صهيب زهير رباع'

    import docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE




    columns_removal=Document()

    #columns_removal.add_paragraph(book_to).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L1=columns_removal.add_paragraph()
    L1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner=L1.add_run(book_to)
    runner.bold=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True



    #columns_removal.add_paragraph(book_subject).alignment=WD_ALIGN_PARAGRAPH.RIGHT
    L2=columns_removal.add_paragraph()
    L2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner=L2.add_run(book_subject)
    # runner.font.BoldBi=True
    runner.bold=True
    runner.underline=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True


    #columns_removal.add_paragraph(book_paragraph1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    L3=columns_removal.add_paragraph()
    L3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L3.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    L3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    runner=L3.add_run(book_paragraph1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    runner.font.rtl=True




    #columns_removal.add_paragraph(book_paragraph2).alignment = WD_ALIGN_PARAGRAPH.LEFT
    L4=columns_removal.add_paragraph()
    L4.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    L4.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    runner=L4.add_run(book_paragraph2)
    runner2=L4.add_run()
    runner.font.rtl=True
    runner.font.name="Arial"
    runner.font.size=Pt(15)



    #columns_removal.add_paragraph(book_s1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L5=columns_removal.add_paragraph()
    L5.alignment=WD_ALIGN_PARAGRAPH.CENTER
    runner=L5.add_run(book_s1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    # runner.font.rtl=True


    #columns_removal.add_paragraph(book_s2).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    L6=columns_removal.add_paragraph()
    L6.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    runner=L6.add_run(book_s2)
    # runner.font.rtl=True
    runner.bold=True
    runner.font.name='Arial'
    runner.font.size=Pt(15)



    #columns_removal.add_paragraph(book_s3).alignment = WD_ALIGN_PARAGRAPH.LEFT
    L7=columns_removal.add_paragraph()
    L7.alignment=WD_ALIGN_PARAGRAPH.LEFT
    runner=L7.add_run(book_s3)
    # runner.font.rtl=True
    runner.font.name='Arial'
    runner.font.size=Pt(14)
    #-----------------------------------------------------------
    L8=columns_removal.add_paragraph()
    L8.alignment=WD_ALIGN_PARAGRAPH.LEFT
    runner=L8.add_run(book_s4)
    # runner.font.rtl=True
    runner.font.name='Arial'
    runner.font.size=Pt(14)
    #----------------------------------------------------
    L9=columns_removal.add_paragraph()
    L9.alignment=WD_ALIGN_PARAGRAPH.LEFT
    runner=L9.add_run(book_note)
    runner.font.name='Arial'
    runner.font.size=Pt(13)
    runner.underline=True
#     docName="projectOne//static//"+ citizen+".docx"
    # docName=citizen+".docx"
    docName="projectOne//static//"+ citizen+".docx"
    # docName="projectOne\\static\\"+ citizen+".docx"
    columns_removal.save(docName)


    return citizen
