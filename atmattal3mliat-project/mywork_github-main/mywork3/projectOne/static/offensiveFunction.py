def offensiveRemoval(road_length,citizen,citizen_request,land_number,basin_name,basin_number,area,offensive):

    line1='   اشارة الى الطلب  المقدم من المواطن {one} '.format(one=citizen)
    line2= ' رقم ({two}) '.format(two=citizen_request)
    line3='و المتضمن فتح الطريق المرسم على مخططات الاراضي والمساحة بطول ({three}) '.format(three=road_length)
    line4= ' والواصل لقطعة الارض رقم ({four}) ' .format(four=land_number)
    line5= ' من حوض {five} '.format(five=basin_name)
    line6= ' رقم ({six}) '.format(six=basin_number)
    line7= ' في منطقة {seven} '.format(seven=area)

    line8= "   ارجو التكرم بالإيعاز لمن يلزم بازالة العوائق ({eight}) ".format(eight=offensive)
    line9= "الواقعة ضمن حرم الطريق اعلاه وذلك لاعتراضها لاعمال الفتوح والتسويـــة للطريق من قبل آليات مديرية اشغال محافظة مادبا."
    import docx
    from docx import Document
    book_to= '\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'\n'+'عطوفة محافظ مادبا'
    book_subject="الموضوع : ازالة اعتداء "
    book_paragraph1=line1+line2+line3+line4+line5+line6+line7+'.'
    book_paragraph2=line8+line9
    book_s1='وتفضلوا عطوفتكم قبول فائق الاحترام'
    book_s2= 'مدير اشغال محافظة مادبا ' + '\n' + ' المهندسة عروبه عويس  '
    book_s3=('نسخه / مساعد المدير لشؤون الطــرق'
    +'\n'+'نسخه / ر.ق  قسم الطرق الزراعيــــة'+'\n'+'نسخه / للملــــــــــــــــــــــــــــــــــــف'
    +'\n'+': المرفقات  '+'\n'+ 'صوره عن الاستدعــــــــــاء ' + '\n' +'صورة عن تقرير المســـــاح '+'\n'+
    'صورة عن محطط الاراضي ')
    citizen_name='طريق'+' ' +citizen+'.docx'


    import docx
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.text import WD_LINE_SPACING
    from docx.shared import Pt
    from docx.enum.style import WD_STYLE_TYPE

#--------------------------------------------------------------------------------------
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    sealcoat_text=Document()

#

#-----------------------------------------------------------------------
    # sealcoat_text.add_paragraph(book_to).alignment = WD_ALIGN_PARAGRAPH.CENTER
    #columns_removal.add_paragraph(book_to).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L1=sealcoat_text.add_paragraph()
    L1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    runner=L1.add_run(book_to)
    runner.bold=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True
#-----------------------------------------------------------------------------


    #columns_removal.add_paragraph(book_subject).alignment=WD_ALIGN_PARAGRAPH.RIGHT
#     sealcoat_text.add_paragraph(book_subject).alignment=WD_ALIGN_PARAGRAPH.RIGHT

    L2=sealcoat_text.add_paragraph()
    L2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    runner=L2.add_run(book_subject)
    # runner.font.BoldBi=True
    runner.bold=True
    runner.underline=True
    runner.font.size=Pt(15)
    runner.font.name='Arial'
    # runner.font.rtl=True



#-----------------------------------------------------------------------------------


#sealcoat_text.add_paragraph(book_paragraph1).alignment = WD_ALIGN_PARAGRAPH.LEFT

 #columns_removal.add_paragraph(book_paragraph1).alignment = WD_ALIGN_PARAGRAPH.LEFT
    L3=sealcoat_text.add_paragraph()
    L3.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L3.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    L3.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    runner=L3.add_run(book_paragraph1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    runner.font.rtl=True

#------------------------------------------------------------------------------------

# sealcoat_text.add_paragraph(book_paragraph2).alignment = WD_ALIGN_PARAGRAPH.LEFT

    L4=sealcoat_text.add_paragraph()
    L4.alignment=WD_ALIGN_PARAGRAPH.LEFT
    L4.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY_LOW
    L4.line_spacing=WD_LINE_SPACING.ONE_POINT_FIVE
    runner=L4.add_run(book_paragraph2)
    runner2=L4.add_run()
    runner.font.rtl=True
    runner.font.name="Arial"
    runner.font.size=Pt(15)


#------------------------------------------------------------------------------------
 #columns_removal.add_paragraph(book_s1).alignment = WD_ALIGN_PARAGRAPH.CENTER
 #sealcoat_text.add_paragraph(book_s1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    L5=sealcoat_text.add_paragraph()
    L5.alignment=WD_ALIGN_PARAGRAPH.CENTER
    runner=L5.add_run(book_s1)
    runner.font.name='Arial'
    runner.font.size=Pt(15)
    # runner.font.rtl=True

#------------------------------------------------------------------------------------
#sealcoat_text.add_paragraph(book_s2).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    L6=sealcoat_text.add_paragraph()
    L6.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    runner=L6.add_run(book_s2)
    # runner.font.rtl=True
    runner.bold=True
    runner.font.name='Arial'
    runner.font.size=Pt(15)


#------------------------------------------------------------------------------------


# sealcoat_text.add_paragraph(book_s3).alignment = WD_ALIGN_PARAGRAPH.LEFT


    L7=sealcoat_text.add_paragraph()
    L7.alignment=WD_ALIGN_PARAGRAPH.LEFT
    runner=L7.add_run(book_s3)
    # runner.font.rtl=True
    runner.font.name='Arial'
    runner.font.size=Pt(14)

#------------------------------------------------------------------------------------
    citizen_name="projectOne//static//"+ citizen+".docx"
    sealcoat_text.save(citizen_name)
    return citizen
